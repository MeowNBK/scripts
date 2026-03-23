import os
import re
import sys
import gc
import argparse
import concurrent.futures
import docx
from docx.document import Document
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ==========================================
# 0. REGEX & CẤU TRÚC MA TRẬN
# ==========================================
RE_EQ_MARKER = re.compile(r'==\s*')
RE_A_D = re.compile(r'\b([A-Da-d])\s+\.')
RE_A_D_NOSPACE = re.compile(r'\b([A-Da-d]\.)([a-zA-Z])')
RE_MULTI_SPACE = re.compile(r' {2,}')
RE_OPT_START = re.compile(r'(?:^|\s)([A-Da-d]\.)(?=\s|$)')
RE_DIAL_START = re.compile(r'(?:^|\s)([a-d]\.\s*[A-Z][a-z]+:|[A-Z][a-z]{2,}:)')
RE_Q_PREFIX = re.compile(r'^\s*(?:Q[a-z]*|Câu|C|Ex[a-z]*|Bài)?\s*(\d+)[\.\:\-\)]+\s*', re.IGNORECASE)

MATRIX_BOUNDARIES = {2, 4, 16, 19, 24, 30, 35, 40}

FILE_CONFIG = {
    ("PHAT_AM", "0_DEFAULT"): ("1_Pronunciation", "CHUYÊN ĐỀ: PHÁT ÂM (PRONUNCIATION)\n(cùng Mèo mất não 🐧)", "Mark the letter A, B, C, or D to indicate the word whose underlined part differs from the other three in pronunciation."),
    ("GIAO_TIEP", "0_DEFAULT"): ("2_Communication", "CHUYÊN ĐỀ: GIAO TIẾP (COMMUNICATION)\n(cùng Mèo mất não 🐧)", "Mark the letter A, B, C, or D to indicate the correct arrangement of the sentences to make a meaningful dialogue or the best response."),
    ("TU_VUNG", "0_DEFAULT"): ("3_Vocabulary_Grammar", "CHUYÊN ĐỀ: TỪ VỰNG & NGỮ PHÁP (VOCABULARY & GRAMMAR)\n(cùng Mèo mất não 🐧)", "Mark the letter A, B, C, or D to indicate the correct answer to each of the following questions."),
    ("DOC_HIEU", "1_ADVERTISEMENT"): ("4_Reading_1_Advertisement", "CHUYÊN ĐỀ: ĐỌC HIỂU (PHẦN 1 - QUẢNG CÁO/THÔNG BÁO)\n(cùng Mèo mất não 🐧)", "Read the following advertisements/announcements and mark the letter A, B, C, or D to indicate the correct option."),
    ("DOC_HIEU", "2_CLOZE"): ("5_Reading_2_Cloze_Test", "CHUYÊN ĐỀ: ĐỌC HIỂU (PHẦN 2 - ĐIỀN TỪ)\n(cùng Mèo mất não 🐧)", "Read the following passages and mark the letter A, B, C, or D to indicate the correct word or phrase that best fits each of the numbered blanks."),
    ("DOC_HIEU", "3_COMPREHENSION"): ("6_Reading_3_Comprehension", "CHUYÊN ĐỀ: ĐỌC HIỂU (PHẦN 3 - TRẢ LỜI CÂU HỎI)\n(cùng Mèo mất não 🐧)", "Read the following passages and mark the letter A, B, C, or D to indicate the correct answer to each of the questions."),
    ("VIET", "1_TRANSFORMATION"): ("7_Writing_1_Transformation", "CHUYÊN ĐỀ: VIẾT (PHẦN 1 - VIẾT LẠI CÂU)\n(cùng Mèo mất não 🐧)", "Mark the letter A, B, C, or D to indicate the sentence that is closest in meaning to each of the following sentences."),
    ("VIET", "2_COMBINATION"): ("8_Writing_2_Combination", "CHUYÊN ĐỀ: VIẾT (PHẦN 2 - GHÉP CÂU)\n(cùng Mèo mất não 🐧)", "Mark the letter A, B, C, or D to indicate the sentence that best combines each pair of sentences in the following questions.")
}

# ==========================================
# 1. BỘ PHÂN TÍCH CẤU TRÚC (AST NODE)
# ==========================================
class CleanRun:
    def __init__(self, run):
        self.text = run.text
        self.bold = run.bold
        self.italic = run.italic
        self.underline = run.underline
        self.superscript = run.font.superscript
        self.subscript = run.font.subscript
        self.color_rgb = str(run.font.color.rgb) if run.font.color and run.font.color.rgb else None
        try: self.highlight_color = run.font.highlight_color
        except: self.highlight_color = None

class CleanParagraph:
    def __init__(self, p):
        self.text = p.text
        self.runs = [CleanRun(r) for r in p.runs if r.text]
        try: self.left_indent = p.paragraph_format.left_indent
        except: self.left_indent = None

class QuestionNode:
    def __init__(self, raw_id=None):
        self.paragraphs = []
        self.raw_id = raw_id
        
    def get_hash(self):
        stem_text = ""
        options = []
        for p in self.paragraphs:
            t = p.text.strip()
            t = RE_Q_PREFIX.sub('', t)
            parts = re.split(r'\b[A-Da-d]\.', t)
            if len(parts) > 1:
                stem_text += parts[0]
                options.extend([opt.strip() for opt in parts[1:] if opt.strip()])
            else:
                stem_text += t
        clean_stem = re.sub(r'\W+', '', stem_text).lower()
        clean_opts = sorted([re.sub(r'\W+', '', opt).lower() for opt in options])
        return clean_stem + "".join(clean_opts)

class SectionBlock:
    def __init__(self, topic, sub_topic="0_DEFAULT"):
        self.topic = topic
        self.sub_topic = sub_topic
        self.context = []   
        self.questions = [] 
        
    def get_hash(self):
        ctx_text = "".join([p.text for p in self.context])
        qs_text = "".join([q.get_hash() for q in self.questions])
        return re.sub(r'\W+', '', ctx_text + qs_text).lower()

def is_answer_style(run):
    if not run or run == "STEM_PREFIX": return False
    try:
        if run.highlight_color and run.highlight_color != 0: return True
    except: pass
    if run.color_rgb and len(run.color_rgb) == 6:
        try:
            r, g, b = int(run.color_rgb[0:2], 16), int(run.color_rgb[2:4], 16), int(run.color_rgb[4:6], 16)
            if (r > 150 and g < 100 and b < 100) or \
               (r > 150 and g > 150 and b < 100) or \
               (r < 100 and g > 150 and b < 100):
                return True
        except: pass
    return False

# ==========================================
# 2. LÕI RENDERING VIRTUAL DOM
# ==========================================
class CharFmt:
    __slots__ = ['char', 'run', 'is_answer']
    def __init__(self, char, run, is_answer=False):
        self.char = char
        self.run = run
        self.is_answer = is_answer

def extract_char_fmts(p):
    char_fmts = []
    for r in p.runs:
        is_ans = is_answer_style(r)
        for c in r.text:
            # FIX: Quét sạch mọi ký tự ngắt trang ẩn, ngắt dòng ẩn gây khoảng trắng dị thường
            if c in ['\n', '\r', '\t', '\xa0', '\v', '\f', '\x0c']: c = ' '
            char_fmts.append(CharFmt(c, r, is_ans))
            
    text = "".join([cf.char for cf in char_fmts])
    for m in reversed(list(RE_EQ_MARKER.finditer(text))):
        start, end = m.start(), m.end()
        if end < len(char_fmts): char_fmts[end].is_answer = True 
        del char_fmts[start:end]

    text = "".join([cf.char for cf in char_fmts])
    for m in reversed(list(RE_A_D.finditer(text))):
        del char_fmts[m.start(1)+1 : m.end()-1]
        
    text = "".join([cf.char for cf in char_fmts])
    for m in reversed(list(RE_A_D_NOSPACE.finditer(text))):
        char_fmts.insert(m.start(2), CharFmt(' ', None))
        
    text = "".join([cf.char for cf in char_fmts])
    for m in reversed(list(RE_MULTI_SPACE.finditer(text))):
        del char_fmts[m.start()+1 : m.end()]
        
    return char_fmts

def apply_replace_map(char_fmts, replace_map):
    for old_n, new_n in sorted(replace_map.items(), key=lambda x: -x[0]):
        text = "".join([cf.char for cf in char_fmts])
        pattern = rf'([\(\[_]+)\s*{old_n}\s*([\)\]_]+)'
        for m in reversed(list(re.finditer(pattern, text))):
            replacement = m.group(1) + str(new_n) + m.group(2)
            start, end = m.start(), m.end()
            run_ref = char_fmts[start].run
            del char_fmts[start:end]
            for i, c in enumerate(replacement):
                char_fmts.insert(start + i, CharFmt(c, run_ref))

def split_char_fmts(char_fmts):
    text = "".join([cf.char for cf in char_fmts])
    split_targets = set([0])
    
    for m in RE_OPT_START.finditer(text): split_targets.add(m.start(1))
    for m in RE_DIAL_START.finditer(text): split_targets.add(m.start(1))
        
    targets = sorted(list(split_targets))
    targets.append(len(text))
    
    chunks = []
    for i in range(len(targets)-1):
        chunk = char_fmts[targets[i]:targets[i+1]]
        while chunk and chunk[0].char == ' ': chunk.pop(0)
        while chunk and chunk[-1].char == ' ': chunk.pop()
        
        if chunk:
            if any(cf.is_answer for cf in chunk):
                for cf in chunk: cf.is_answer = True
            chunks.append(chunk)
    return chunks

def setup_exam_format(p, is_stem=False, is_context=False, is_dialogue=False, old_left_indent=None):
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(0)
    p.paragraph_format.first_line_indent = Cm(0)
    
    # FIX: Chốt cứng không cho Word tự động nhảy trang lung tung
    p.paragraph_format.page_break_before = False
    p.paragraph_format.keep_with_next = False
    p.paragraph_format.keep_together = False
    
    if is_context:
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        if old_left_indent: p.paragraph_format.left_indent = old_left_indent
    elif is_dialogue:
        p.paragraph_format.line_spacing = 1.0 
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(0.5)
    else: 
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_before = Pt(6) 
        p.paragraph_format.space_after = Pt(2)

def _add_run(p, text, old_run, is_answer=False):
    if not text: return
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    
    if old_run == "STEM_PREFIX": r.bold = True
    elif old_run is not None:
        r.font.superscript = old_run.superscript
        r.font.subscript = old_run.subscript
        if is_answer:
            r.font.color.rgb = RGBColor(255, 0, 0)
            r.bold = True
            r.underline = True
        else:
            if old_run.color_rgb:
                try: r.font.color.rgb = RGBColor(int(old_run.color_rgb[0:2],16), int(old_run.color_rgb[2:4],16), int(old_run.color_rgb[4:6],16))
                except: pass
            r.bold = old_run.bold
            r.italic = old_run.italic
            r.underline = old_run.underline
    elif is_answer:
        r.font.color.rgb = RGBColor(255, 0, 0)
        r.bold = True
        r.underline = True

def write_chunk_to_paragraph(p, chunk):
    current_run_ref = -1
    current_text = ""
    is_transparent = False
    current_is_ans = False

    for cf in chunk:
        char_transparent = cf.char in ['\t', '\n']
        if cf.run is not current_run_ref or char_transparent != is_transparent or cf.is_answer != current_is_ans:
            if current_text: _add_run(p, current_text, None if is_transparent else current_run_ref, current_is_ans)
            current_run_ref = cf.run
            current_text = cf.char
            is_transparent = char_transparent
            current_is_ans = cf.is_answer
        else:
            current_text += cf.char
            
    if current_text: _add_run(p, current_text, None if is_transparent else current_run_ref, current_is_ans)

def render_options_table(doc, opt_chunks):
    if not opt_chunks: return
    max_len = max(len("".join([cf.char for cf in chunk])) for chunk in opt_chunks)
    num_opts = len(opt_chunks)
    
    if num_opts == 4:
        if max_len < 26: rows, cols = 1, 4    
        elif max_len < 50: rows, cols = 2, 2  
        else: rows, cols = 4, 1               
    elif num_opts == 3:
        if max_len < 30: rows, cols = 1, 3
        else: rows, cols = 3, 1
    elif num_opts == 2:
        if max_len < 50: rows, cols = 1, 2
        else: rows, cols = 2, 1
    else:
        rows, cols = num_opts, 1
        
    table = doc.add_table(rows=rows, cols=cols)
    table.autofit = False
    
    total_width = Cm(17.0)
    col_width = total_width / cols
    
    for row in table.rows:
        row.height_rule = docx.enum.table.WD_ROW_HEIGHT_RULE.AUTO
        for cell in row.cells: cell.width = col_width
            
    for i, chunk in enumerate(opt_chunks):
        r = i // cols
        c = i % cols
        cell = table.cell(r, c)
        p = cell.paragraphs[0]
        # Xóa định dạng nhảy trang rác trong table
        p.paragraph_format.page_break_before = False
        p.paragraph_format.keep_with_next = False
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Cm(0.25)
        p.paragraph_format.first_line_indent = Cm(-0.25) 
        write_chunk_to_paragraph(p, chunk)

def process_context_paragraph(doc, p, replace_map=None):
    char_fmts = extract_char_fmts(p)
    if replace_map: apply_replace_map(char_fmts, replace_map)
    while char_fmts and char_fmts[0].char == ' ': char_fmts.pop(0)
    while char_fmts and char_fmts[-1].char == ' ': char_fmts.pop()
    if char_fmts:
        new_p = doc.add_paragraph(style='Normal')
        setup_exam_format(new_p, is_context=True, old_left_indent=p.left_indent)
        write_chunk_to_paragraph(new_p, char_fmts)

def process_question_node(doc, q_node, stem_prefix=None, replace_map=None):
    char_fmts = []
    for p in q_node.paragraphs:
        char_fmts.extend(extract_char_fmts(p))
        if char_fmts and char_fmts[-1].char != ' ':
            char_fmts.append(CharFmt(' ', None))
            
    if replace_map: apply_replace_map(char_fmts, replace_map)
        
    chunks = split_char_fmts(char_fmts)
    if not chunks: return
    
    first_text = "".join([cf.char for cf in chunks[0]])
    match = RE_Q_PREFIX.match(first_text)
    if match: del chunks[0][:match.end()]
        
    def is_opt(t):
        return bool(re.match(r'^\s*(?:(?:Question\s*\d+|Câu\s*\d+|Ex\s*\d+|Q\s*\d+|\d+)[\.\:\-]\s*)?[A-Da-d]\.', t, flags=re.IGNORECASE))
        
    idx = 0
    while idx < len(chunks) and not "".join([cf.char for cf in chunks[idx]]).strip():
        idx += 1

    # Xử lý thông minh: Câu hỏi không có lời dẫn (Chỉ in Q1 ở dòng trên, đáp án bảng ở dòng dưới)
    if stem_prefix:
        if idx < len(chunks):
            first_real_text = "".join([cf.char for cf in chunks[idx]]).strip()
            if is_opt(first_real_text):
                # Tách Q1 ra một đoạn nhỏ bên trên để bảo vệ khung Grid của bảng A B C D
                p = doc.add_paragraph(style='Normal')
                setup_exam_format(p, is_stem=True)
                # Thu hẹp khoảng cách với bảng đáp án
                p.paragraph_format.space_after = Pt(0)
                _add_run(p, stem_prefix.strip(), "STEM_PREFIX")
            else:
                for c in reversed(stem_prefix): chunks[idx].insert(0, CharFmt(c, "STEM_PREFIX"))
        else:
            p = doc.add_paragraph(style='Normal')
            setup_exam_format(p, is_stem=True)
            _add_run(p, stem_prefix.strip(), "STEM_PREFIX")
            
    while idx < len(chunks):
        chunk = chunks[idx]
        text = "".join([cf.char for cf in chunk]).strip()
        
        if not text:
            idx += 1; continue
            
        if is_opt(text):
            opt_chunks = []
            while idx < len(chunks):
                n_text = "".join([cf.char for cf in chunks[idx]]).strip()
                if is_opt(n_text):
                    opt_chunks.append(chunks[idx])
                    idx += 1
                else: break
            render_options_table(doc, opt_chunks)
        else:
            p = doc.add_paragraph(style='Normal')
            is_dial = bool(RE_DIAL_START.match(text))
            setup_exam_format(p, is_stem=not is_dial, is_dialogue=is_dial)
            write_chunk_to_paragraph(p, chunk)
            idx += 1

# ==========================================
# 3. ĐỊNH TUYẾN THEO MA TRẬN 40 CÂU
# ==========================================
def is_garbage_line(text):
    t = text.strip()
    if not t: return True
    clean_t = re.sub(r'[^a-zA-Z0-9]', '', t).upper()
    if clean_t in ["THEEND", "HET", "END"]: return True
    if "HẾT" in t.upper() and len(t) < 20: return True
    if re.match(r'^(trang|page)\s*\d+.*', t, re.IGNORECASE): return True
    if re.match(r'^(SỞ GIÁO DỤC|ĐỀ THAM KHẢO|KỲ THI|Họ và tên|Thời gian|Số báo danh)', t, re.IGNORECASE): return True
    if t.isupper() and len(t) < 40 and any(x in t for x in ["A. PHONETICS", "B. READING", "C. WRITING", "D. VOCABULARY", "LANGUAGE"]): return True
    return False

def is_instruction_line(text):
    t = text.strip().lower()
    t_stripped = RE_Q_PREFIX.sub('', t)
    if re.match(r'^(part|section|phần)\s+[ivx0-9]+', t): return True
    if re.match(r'^[ivx]+\.\s+[a-z]', t): return True
    if re.match(r'^(i|v|x|\d+)*[\.\:\)]*\s*(mark|choose|read|indicate|circle|complete)\b', t_stripped): return True
    if any(x in t_stripped for x in ["mark the letter", "correct answer", "indicate the word", "closest in meaning", "best combines", "following passage"]): return True
    return False

def get_matrix_mapping(q_id):
    if 1 <= q_id <= 2: return ("PHAT_AM", "0_DEFAULT")
    if 3 <= q_id <= 4: return ("GIAO_TIEP", "0_DEFAULT")
    if 5 <= q_id <= 16: return ("TU_VUNG", "0_DEFAULT")
    if 17 <= q_id <= 19: return ("DOC_HIEU", "1_ADVERTISEMENT")
    if 20 <= q_id <= 24: return ("DOC_HIEU", "2_CLOZE")
    if 25 <= q_id <= 30: return ("DOC_HIEU", "3_COMPREHENSION")
    if 31 <= q_id <= 35: return ("VIET", "1_TRANSFORMATION")
    if 36 <= q_id <= 40: return ("VIET", "2_COMBINATION")
    return None

def create_exam_header(doc, topic_vi):
    header_p = doc.add_paragraph(style='Normal')
    header_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_p.paragraph_format.space_after = Pt(12)
    header_p.add_run("KỲ THI TUYỂN SINH VÀO LỚP 10 THPT\n").bold = True
    header_p.add_run("TÀI LIỆU ÔN TẬP\n").bold = True
    topic_lines = topic_vi.split('\n')
    header_p.add_run(f"{topic_lines[0]}\n").bold = True
    if len(topic_lines) > 1:
        r_meo = header_p.add_run(f"{topic_lines[1]}\n")
        r_meo.italic = True
        r_meo.font.size = Pt(11)
    header_p.add_run("------------------------")
    for r in header_p.runs: r.font.name = 'Times New Roman'; r.font.size = Pt(13) if not r.font.size else r.font.size

def iter_block_items(parent):
    if isinstance(parent, Document): parent_elm = parent.element.body
    elif isinstance(parent, _Cell): parent_elm = parent._tc
    else: return
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P): yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl): yield Table(child, parent)

def parse_document(filepath):
    try: doc = docx.Document(filepath)
    except Exception as e: 
        print(f"[-] Bỏ qua (Lỗi File): {os.path.basename(filepath)}")
        return []

    blocks_dict = {key: SectionBlock(*key) for key in FILE_CONFIG.keys()}
    current_q = None
    last_seen_q_id = 0

    def process_extracted_paragraph(p):
        nonlocal current_q, last_seen_q_id, blocks_dict
        text = p.text.strip()
        if is_garbage_line(text): return
        
        if is_instruction_line(text):
            current_q = None
            return
            
        match = RE_Q_PREFIX.match(text)
        if match:
            q_id = int(match.group(1))
            mapping = get_matrix_mapping(q_id)
            if mapping:
                current_q = QuestionNode(raw_id=q_id)
                current_q.paragraphs.append(CleanParagraph(p))
                blocks_dict[mapping].questions.append(current_q)
                last_seen_q_id = q_id
            else:
                current_q = QuestionNode(raw_id=q_id)
                current_q.paragraphs.append(CleanParagraph(p))
                blocks_dict[("VIET", "2_COMBINATION")].questions.append(current_q)
            return

        if current_q:
            if current_q.raw_id in MATRIX_BOUNDARIES:
                full_text = "\n".join([pa.text for pa in current_q.paragraphs])
                if re.search(r'\b[Dd]\.', full_text):
                    if text and not re.match(r'^[A-Da-d]\.', text):
                        if text[0].isupper() or text.startswith('['):
                            current_q = None 

        if current_q:
            current_q.paragraphs.append(CleanParagraph(p))
        else:
            anticipated_q_id = last_seen_q_id + 1
            mapping = get_matrix_mapping(anticipated_q_id)
            if mapping:
                blocks_dict[mapping].context.append(CleanParagraph(p))

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph): process_extracted_paragraph(block)
        elif isinstance(block, Table):
            for row in block.rows:
                for cell in row.cells:
                    for p in cell.paragraphs: process_extracted_paragraph(p)

    return list(blocks_dict.values())

# ==========================================
# 4. ENGINE BUILD & ĐIỀU PHỐI LUỒNG
# ==========================================
def build_chuyen_de(all_blocks, out_dir, output_suffix=""):
    os.makedirs(out_dir, exist_ok=True)
    block_map = {k: [] for k in FILE_CONFIG.keys()}
    for b in all_blocks:
        key = (b.topic, b.sub_topic)
        if key in block_map and (b.questions or b.context): block_map[key].append(b)

    for key, blocks in block_map.items():
        if not blocks: continue
        
        topic_id, sub_topic_id = key
        file_name, topic_vi, instruction = FILE_CONFIG[key]
        
        doc = docx.Document()
        for section in doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)

        create_exam_header(doc, topic_vi)
        
        # FIX: IN HƯỚNG DẪN 1 LẦN DUY NHẤT TRÊN ĐẦU FILE
        inst_p = doc.add_paragraph(style='Normal')
        inst_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        inst_p.paragraph_format.space_before = Pt(12)
        inst_p.paragraph_format.space_after = Pt(12)
        r_inst = inst_p.add_run(instruction)
        r_inst.bold = True
        r_inst.italic = True
        r_inst.font.name = 'Times New Roman'
        r_inst.font.size = Pt(12)

        seen_hashes = set()
        global_ex_idx = 1
        global_q_idx = 1

        for b in blocks:
            b_hash = b.get_hash()
            if b_hash in seen_hashes: continue
            seen_hashes.add(b_hash)

            valid_context = [x for x in b.context if x.text.strip()]
            valid_questions = b.questions
            
            if not valid_context and not valid_questions: continue

            if topic_id == "DOC_HIEU":
                # FIX: Chỉ in chữ Ex 1. gọn gàng cho phần Reading
                p_ex = doc.add_paragraph(style='Normal')
                p_ex.paragraph_format.space_before = Pt(18)
                p_ex.paragraph_format.space_after = Pt(6)
                r_ex = p_ex.add_run(f"Ex {global_ex_idx}.")
                r_ex.bold = True
                r_ex.font.name = 'Times New Roman'
                r_ex.font.size = Pt(12)
                global_ex_idx += 1
                
                replace_map = {}
                for idx, q in enumerate(valid_questions):
                    if q.raw_id is not None: replace_map[q.raw_id] = global_q_idx + idx

                for p in valid_context:
                    process_context_paragraph(doc, p, replace_map=replace_map)

                for q in valid_questions:
                    stem_prefix = f"Q{global_q_idx}. " 
                    process_question_node(doc, q, stem_prefix=stem_prefix, replace_map=replace_map)
                    global_q_idx += 1
            else:
                for q in valid_questions:
                    q_hash = q.get_hash()
                    if q_hash in seen_hashes: continue
                    seen_hashes.add(q_hash)

                    stem_prefix = f"Q{global_q_idx}. " 
                    process_question_node(doc, q, stem_prefix=stem_prefix, replace_map=None)
                    global_q_idx += 1 

        end_p = doc.add_paragraph(style='Normal')
        end_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        end_p.paragraph_format.space_before = Pt(24)
        run_het = end_p.add_run("----------- HẾT -----------\n")
        run_het.bold = True
        run_het.font.name = 'Times New Roman'
        run_het.font.size = Pt(12)
        
        out_filename = f"{file_name}{output_suffix}.docx"
        doc.save(os.path.join(out_dir, out_filename))

def process_single_file(filepath):
    parent_dir = os.path.basename(os.path.dirname(filepath))
    blocks = parse_document(filepath)
    return (parent_dir, blocks)

def main():
    parser = argparse.ArgumentParser(description="Engine V7.6")
    parser.add_argument("base_dir", nargs="?", default=".", help="Thư mục gốc chứa các thư mục con")
    args = parser.parse_args()

    base_dir = args.base_dir
    print("[*] Khởi động V7.6 (The Perfect Book) - Dọn rác & Tối ưu Giao diện...")

    all_dirs = sorted([os.path.join(base_dir, d) for d in os.listdir(base_dir) if os.path.isdir(os.path.join(base_dir, d))])
    subdirs = [d for d in all_dirs if re.search(r'\d+-\d+', os.path.basename(d))]

    if not subdirs:
        print("[-] Không tìm thấy dữ liệu đầu vào hợp lệ.")
        sys.exit(0)

    all_files = []
    for subdir in subdirs:
        files = sorted([os.path.join(subdir, f) for f in os.listdir(subdir) if f.endswith('.docx') and not f.startswith('~')])
        all_files.extend(files)

    print(f"[*] Đang càn quét {len(all_files)} files...")

    grouped_blocks = {}
    with concurrent.futures.ProcessPoolExecutor() as executor:
        results = executor.map(process_single_file, all_files)
        for i, res in enumerate(results):
            if res:
                parent_dir, blocks = res
                if parent_dir not in grouped_blocks: grouped_blocks[parent_dir] = []
                grouped_blocks[parent_dir].extend(blocks)
                
            if (i + 1) % 50 == 0 or (i + 1) == len(all_files):
                print(f"  -> Đã cứu hộ & định tuyến xong {i + 1}/{len(all_files)} files...")

    main_out_dir = os.path.join(base_dir, "Output")
    os.makedirs(main_out_dir, exist_ok=True)

    print("[*] Bắt đầu xuất file Word chia theo từng phân khu...")
    
    for parent_dir, blocks in grouped_blocks.items():
        if not blocks: continue
        sub_out_dir = os.path.join(main_out_dir, parent_dir)
        print(f"  -> Đang đóng gói chuyên đề cho: {parent_dir}")
        build_chuyen_de(blocks, sub_out_dir)

    print(f"[+] Hoàn tất tuyệt đối! Tài liệu đã chuẩn như Sách Giáo Khoa, lưu tại: {main_out_dir}")

if __name__ == "__main__":
    main()